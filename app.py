from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt
from agent import resume_agent
from resume_generator import generate_summary
import os

app = Flask(__name__)

OUTPUT_FOLDER = "generated_resumes"

if not os.path.exists(OUTPUT_FOLDER):
    os.makedirs(OUTPUT_FOLDER)


@app.route("/")
def home():
    return render_template("index.html")


@app.route("/preview", methods=["POST"])
def preview():

    data = request.form.to_dict()

    template_choice = request.form.get("template_choice")

    if template_choice == "ai":
        template = resume_agent(data)
    else:
        template = template_choice

    summary = generate_summary(data)

    return render_template(
        "preview.html",
        data=data,
        summary=summary,
        template_name=template
    )


@app.route("/download", methods=["POST"])
def download():

    template = request.form.get("template")

    data = {
        "name": request.form.get("name"),
        "email": request.form.get("email"),
        "phone": request.form.get("phone"),
        "skills": request.form.get("skills"),
        "education": request.form.get("education"),
        "experience": request.form.get("experience")
    }

    summary = request.form.get("summary")

    document = Document()

    # Different fonts for different templates
    font_map = {
        "modern": "Calibri",
        "developer": "Courier New",
        "minimal": "Arial",
        "professional": "Times New Roman",
        "elegant": "Georgia",
        "colorful": "Verdana"
    }

    font_name = font_map.get(template, "Calibri")

    style = document.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(12)

    document.add_heading(data["name"], level=0)

    document.add_paragraph(f"Email: {data['email']}")
    document.add_paragraph(f"Phone: {data['phone']}")

    document.add_heading("Professional Summary", level=1)
    document.add_paragraph(summary)

    document.add_heading("Technical Skills", level=1)
    document.add_paragraph(data["skills"])

    document.add_heading("Education", level=1)
    document.add_paragraph(data["education"])

    document.add_heading("Experience", level=1)
    document.add_paragraph(data["experience"])

    file_path = os.path.join(OUTPUT_FOLDER, "resume.docx")

    document.save(file_path)

    return send_file(file_path, as_attachment=True)


if __name__ == "__main__":
    app.run(debug=True)